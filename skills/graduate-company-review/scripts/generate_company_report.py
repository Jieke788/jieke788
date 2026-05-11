#!/usr/bin/env python3
"""
Generate a formal Chinese Word report for graduate company assessments.

Usage:
  python generate_company_report.py payload.json
"""

from __future__ import annotations

import json
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(38, 38, 38)
GRAY = RGBColor(102, 102, 102)
LIGHT = "D9EAF7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size: int | None = None, bold: bool | None = None, color=None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int) -> None:
    style = "Heading 1" if level == 1 else "Heading 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(6 if level == 1 else 4)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    r = p.add_run(text)
    set_run_font(r, size=14 if level == 1 else 12, bold=True, color=BLUE if level == 1 else DARK)


def add_paragraph(doc: Document, text: str, center: bool = False, color=None) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=11, color=color)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=11)


def add_title_page(doc: Document, payload: dict) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(payload["report_title"])
    set_run_font(r, size=22, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(payload["company_name"])
    set_run_font(r, size=18, bold=True)

    add_paragraph(doc, f"面向对象：{payload['audience']}", center=True, color=GRAY)

    for _ in range(5):
        doc.add_paragraph()

    for line in [
        f"报告结论：{payload['verdict']}",
        f"评估口径：{payload['scope']}",
        f"报告日期：{payload['report_date']}",
    ]:
        add_paragraph(doc, line, center=True)

    doc.add_page_break()


def add_summary_table(doc: Document, payload: dict) -> None:
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("公司", payload["company_name"]),
        ("结论", payload["verdict"]),
        ("适合人群", payload["best_for"]),
        ("核心判断", payload["summary"]),
    ]
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        set_cell_shading(table.cell(i, 0), LIGHT)
        for j in range(2):
            cell = table.cell(i, j)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run_font(r, size=11, bold=(j == 0))
    doc.add_paragraph()


def add_score_table(doc: Document, score_rows: list[dict]) -> None:
    add_heading(doc, "综合评分卡", 2)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["维度", "评分", "说明"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, LIGHT)
        for p in cell.paragraphs:
            for r in p.runs:
                set_run_font(r, size=11, bold=True)

    for row in score_rows:
        cells = table.add_row().cells
        values = [row["dimension"], row["score"], row["note"]]
        for idx, value in enumerate(values):
            cells[idx].text = value
            for p in cells[idx].paragraphs:
                for r in p.runs:
                    set_run_font(r, size=11)
    doc.add_paragraph()


def build_document(payload: dict) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    base = doc.styles["Normal"]
    base.font.name = "Microsoft YaHei"
    base._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    base.font.size = Pt(11)

    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    add_title_page(doc, payload)
    add_heading(doc, "一、执行摘要", 1)
    add_summary_table(doc, payload)

    for section_data in payload["sections"]:
        add_heading(doc, section_data["title"], 1)
        for block in section_data["blocks"]:
            block_type = block["type"]
            if block_type == "heading":
                add_heading(doc, block["text"], 2)
            elif block_type == "paragraph":
                add_paragraph(doc, block["text"])
            elif block_type == "bullet":
                add_bullet(doc, block["text"])
            elif block_type == "score_table":
                add_score_table(doc, payload["score_rows"])
            else:
                raise ValueError(f"Unsupported block type: {block_type}")

    return doc


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: python generate_company_report.py <payload.json>")
        return 1

    payload_path = Path(sys.argv[1])
    payload = json.loads(payload_path.read_text(encoding="utf-8-sig"))

    final_out = Path(payload["output_path"])
    temp_out = Path(payload.get("temp_output_path") or payload["output_path"])

    doc = build_document(payload)
    temp_out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(temp_out))

    if temp_out != final_out:
        final_out.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(str(temp_out), str(final_out))

    print(final_out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
